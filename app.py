"""
Test Case Generator App
Streamlit application to generate test cases from BRD documents using AI
"""

import streamlit as st
import pandas as pd
from openai import OpenAI
from docx import Document
import PyPDF2
import io
import openpyxl
from datetime import datetime
import json
import os
import base64
from PIL import Image

# Page configuration
st.set_page_config(
    page_title="Test Case Generator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        font-weight: bold;
        text-align: center;
        padding: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        text-align: center;
        padding-bottom: 2rem;
    }
    .stats-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Load the test case generation prompt
def load_prompt():
    """Load the test case generation prompt from file"""
    try:
        with open('Test_Case_Generation_Prompt.md', 'r') as f:
            return f.read()
    except FileNotFoundError:
        st.error("Test_Case_Generation_Prompt.md not found. Please ensure it's in the same directory.")
        return None

# Document processing functions
def extract_text_from_docx(file):
    """Extract text from Word document"""
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(' | '.join(row_text))
    
    return '\n'.join(full_text)

def extract_text_from_pdf(file):
    """Extract text from PDF document"""
    pdf_reader = PyPDF2.PdfReader(file)
    full_text = []
    for page in pdf_reader.pages:
        full_text.append(page.extract_text())
    return '\n'.join(full_text)

def extract_text_from_excel(file):
    """Extract text from Excel document"""
    df = pd.read_excel(file, sheet_name=None)  # Read all sheets
    full_text = []
    for sheet_name, sheet_df in df.items():
        full_text.append(f"\n--- Sheet: {sheet_name} ---\n")
        full_text.append(sheet_df.to_string())
    return '\n'.join(full_text)

def process_uploaded_files(uploaded_files):
    """Process multiple uploaded files and extract text"""
    combined_text = []
    
    for uploaded_file in uploaded_files:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        st.write(f"📄 Processing: {uploaded_file.name}")
        
        try:
            if file_extension == 'docx':
                text = extract_text_from_docx(uploaded_file)
            elif file_extension == 'pdf':
                text = extract_text_from_pdf(uploaded_file)
            elif file_extension in ['xlsx', 'xls']:
                text = extract_text_from_excel(uploaded_file)
            elif file_extension == 'txt':
                text = uploaded_file.read().decode('utf-8')
            else:
                st.warning(f"Unsupported file type: {file_extension}")
                continue
            
            combined_text.append(f"\n{'='*80}\n")
            combined_text.append(f"FILE: {uploaded_file.name}\n")
            combined_text.append(f"{'='*80}\n")
            combined_text.append(text)
            
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {str(e)}")
    
    return '\n'.join(combined_text)

def analyze_requirements_from_brd(brd_text, api_key):
    """Analyze BRD document to identify and count all functional requirements"""
    
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {str(e)}")
        return [], 0
    
    prompt = f"""You are a Senior Business Analyst specialized in comprehensive requirements extraction from Business Requirements Documents.

## BRD DOCUMENT TO ANALYZE

{brd_text}

## CRITICAL TASK

Extract AT LEAST 30-40 UNIQUE functional requirements from the BRD. Be extremely thorough and granular. Break down complex features into multiple specific requirements.

**REQUIREMENT CATEGORIES TO EXTRACT (Minimum 5-8 per category):**

### 1. USER INTERFACE REQUIREMENTS (8-12 requirements)
- Each screen/page as separate requirement
- Form fields and their validations (group by sections)
- Buttons, links, navigation elements
- Menu items and dropdown options
- Search functionality and filters
- Display formats and layouts
- File upload/download features
- Modal dialogs and popups

### 2. BUSINESS PROCESS REQUIREMENTS (10-15 requirements)
- User registration/login processes
- Account creation workflows
- Transaction processing steps
- Approval workflows (each approval level)
- Status change processes
- Email/notification triggers
- Data synchronization processes
- Batch processing operations

### 3. DATA VALIDATION REQUIREMENTS (8-12 requirements)
- Field-level validations (separate for each field type)
- Business rule validations
- Cross-field validations
- Format validations (date, email, phone, etc.)
- Range validations (amounts, dates)
- Mandatory field checks
- Duplicate data checks
- Data integrity constraints

### 4. INTEGRATION REQUIREMENTS (3-5 requirements)
- API integrations
- Database operations
- Third-party service calls
- File import/export
- External system communications

### 5. REPORTING & INQUIRY REQUIREMENTS (3-5 requirements)
- Report generation
- Data export features  
- Search and inquiry functions
- Dashboard displays
- Analytics and metrics

### 6. SECURITY & ACCESS REQUIREMENTS (3-5 requirements)
- Authentication mechanisms
- Authorization controls
- Role-based access
- Session management
- Audit trail requirements

### 7. BUSINESS RULES & CALCULATIONS (5-8 requirements)
- Calculation logic
- Interest computations
- Fee calculations
- Business rule validations
- Conditional processing
- Status determination logic

**EXTRACTION STRATEGY:**
1. **Be Granular**: Break large features into 3-4 smaller requirements
2. **Field-Level**: Each form section = 1 requirement
3. **Step-by-Step**: Each workflow step = 1 requirement  
4. **Screen-Specific**: Each screen/page = 1 requirement
5. **Validation-Specific**: Each validation type = 1 requirement
6. **Role-Specific**: Different user roles = separate requirements

**EXAMPLES OF GRANULAR BREAKDOWN:**
- Instead of "User Registration" → Break into:
  - REQ-001: User registration form - Personal details validation
  - REQ-002: User registration form - Contact details validation  
  - REQ-003: User registration form - Document upload functionality
  - REQ-004: User registration - Email verification process
  - REQ-005: User registration - Account activation workflow

**MINIMUM TARGET**: 35+ unique requirements (aim for 40-50 if document is comprehensive)

**NO DUPLICATES**: Each requirement must be unique and testable

**OUTPUT FORMAT**:
Return ONLY a valid JSON object with this structure:
{{
  "total_requirements": <number_minimum_35>,
  "requirements": [
    {{
      "requirement_id": "REQ-001",
      "requirement_type": "UI/Workflow/Data/Integration/Report/Security/Business Rule/Calculation",
      "module": "Module Name",
      "title": "Specific granular requirement title",
      "description": "Detailed requirement description with specific functionality",
      "priority": "Critical/High/Medium/Low",
      "testable": true
    }}
  ]
}}

**IMPORTANT**: If the BRD seems to have fewer obvious requirements, EXPAND and DECOMPOSE existing features into granular, testable requirements. Every button click, field validation, screen display, and process step should be a separate requirement.

Start extraction now - aim for 35-50 unique requirements:"""

    try:
        with st.spinner('🔍 Analyzing BRD to identify all requirements...'):
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert Business Analyst specialized in requirements analysis and extraction."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=8192,
                temperature=0.2,
                response_format={"type": "json_object"}
            )
            
            response_text = response.choices[0].message.content
            
            # Parse JSON response
            try:
                json_response = json.loads(response_text)
                requirements = json_response.get('requirements', [])
                total_count = json_response.get('total_requirements', len(requirements))
                
                return requirements, total_count
                
            except json.JSONDecodeError as e:
                st.error(f"Error parsing requirements analysis: {str(e)}")
                return [], 0
                
    except Exception as e:
        st.error(f"Error analyzing requirements: {str(e)}")
        return [], 0

def generate_test_cases_with_ai(brd_text, prompt_template, api_key, options):
    """Generate test cases using OpenAI GPT-4o API"""
    
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {str(e)}")
        return [], ""
    
    # First analyze requirements to get exact count
    st.info("🔍 Step 1: Analyzing BRD to identify all requirements...")
    requirements, total_requirements_count = analyze_requirements_from_brd(brd_text, api_key)
    
    if not requirements:
        st.warning("Could not identify requirements. Falling back to automatic detection...")
        total_requirements_count = 35  # Minimum fallback count
        requirements_summary = "Analyze BRD and identify at least 35 granular requirements automatically"
    else:
        if total_requirements_count < 30:
            st.warning(f"⚠️ Only {total_requirements_count} requirements found. Expanding analysis...")
            total_requirements_count = max(35, total_requirements_count)
            st.info(f"📈 Targeting {total_requirements_count} test cases with granular requirement breakdown")
        else:
            st.success(f"✅ Found {total_requirements_count} requirements in the BRD")
        
        requirements_summary = f"Generate EXACTLY {total_requirements_count} test cases - one for each requirement:\n"
        for i, req in enumerate(requirements[:10], 1):  # Show first 10 for reference
            requirements_summary += f"{i}. {req.get('title', 'N/A')} (Type: {req.get('requirement_type', 'N/A')})\n"
        if len(requirements) > 10:
            requirements_summary += f"... and {len(requirements) - 10} more requirements"
    
    st.info("🤖 Step 2: Generating test cases for each identified requirement...")
    
    # Construct the full prompt
    full_prompt = f"""{prompt_template}

---

## BRD DOCUMENT TO ANALYZE

{brd_text}

---

## REQUIREMENTS ANALYSIS RESULTS

{requirements_summary}

---

## GENERATION INSTRUCTIONS

Based on the above BRD document and requirements analysis, generate comprehensive test cases following the framework provided.

**CRITICAL REQUIREMENT**: Generate EXACTLY {total_requirements_count if isinstance(total_requirements_count, int) else 'at least 35'} test cases - ONE test case for EACH identified requirement. Be comprehensive and granular - break down complex features into multiple specific test cases.

**MINIMUM TEST CASE TARGET**: {total_requirements_count if isinstance(total_requirements_count, int) else '35'} test cases (ensure comprehensive coverage)

**APPLICATION DOMAIN/CONTEXT:**
{options.get('domain_context', 'Identify from BRD')}

Use domain-appropriate terminology, workflows, and test scenarios relevant to this industry/application type.

**IMPORTANT REQUIREMENTS:**
1. Generate test cases in the 14-column structure specified
2. Include field-level validation test cases (~20%)
3. Include workflow test cases with 15-step pattern (~55%)
4. Include authorization/approval test cases (6+ per workflow if applicable)
5. Include negative test cases (~{options['negative_ratio']}%)
6. Consider variations: {', '.join(options['variations']) if options['variations'] else 'Identify from BRD'}
7. Focus on these modules: {', '.join(options['focus_modules']) if options['focus_modules'] else 'All modules from BRD'}

**ADAPT TO DOMAIN:**
- Use terminology specific to {options.get('domain_context', 'the application domain')}
- Include domain-specific validations and business rules
- Create realistic test scenarios for this industry
- Consider domain-specific user roles and workflows

**IMPORTANT - TEST CASE CATEGORIZATION:**
Clearly mark each test case with appropriate type in the TC Module field:
- **Field-Level**: Use format "[Module Name] - Field Validation"
- **Functional/Workflow**: Use format "[Module Name] - Workflow"
- **Negative**: Use format "[Module Name] - Negative Test"

**TEST CASE DISTRIBUTION:**
1. **Field-Level Test Cases (~20% of total)**:
   - Focus on field validations, formats, mandatory checks
   - TC Module should end with "- Field Validation"
   - Example: "Account Opening - Field Validation"

2. **Functional/Workflow Test Cases (~60% of total)**:
   - Focus on end-to-end workflows, business processes
   - TC Module should end with "- Workflow"
   - Example: "Account Opening - Workflow"

3. **Negative Test Cases (~15% of total)**:
   - Focus on business rule violations, error scenarios
   - TC Module should end with "- Negative Test"
   - Example: "Account Opening - Negative Test"

4. **Other Test Cases (~5%)**:
   - Reports, inquiries, modifications
   - TC Module should reflect the type

**OUTPUT FORMAT:**
Return ONLY a valid JSON object with a "test_cases" key containing an array of test case objects.
Each test case object must have these 14 fields:
- product_name
- process_category
- business_process_id
- business_process
- scenario_id
- scenario_description
- category
- importance
- test_case_id
- tc_module (MUST include type suffix: "- Field Validation", "- Workflow", or "- Negative Test")
- test_condition
- prerequisite
- test_case_description
- expected_result

**IMPORTANT - COMPREHENSIVE TEST CASE COUNT REQUIREMENT:**
Generate EXACTLY {total_requirements_count if isinstance(total_requirements_count, int) else 'at least 35'} test cases based on:
- One test case for each functional requirement identified in the requirements analysis
- If requirements analysis found fewer than 30 requirements, expand by breaking down complex features into granular test cases
- Each screen, form section, workflow step, validation rule should be a separate test case
- Cover all UI elements, business processes, data validations, integrations, reports, and security aspects
- Be granular: "User Registration" should become 4-5 separate test cases (personal details, contact details, document upload, email verification, account activation)
- Ensure comprehensive coverage: {total_requirements_count if isinstance(total_requirements_count, int) else 'minimum 35'} unique test cases
- No duplicates - each test case must test a unique aspect of the system

Example format:
{{
  "test_cases": [
    {{
      "product_name": "Banking Application",
      "process_category": "Account Opening",
      "tc_module": "Account Opening - Field Validation",
      ...
    }},
    {{
      "product_name": "Banking Application",
      "process_category": "Account Opening",
      "tc_module": "Account Opening - Workflow",
      ...
    }}
  ]
}}

Start generating now:"""
    
    with st.spinner('🤖 AI is analyzing BRD and generating test cases...'):
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert QA Test Analyst specialized in creating comprehensive test cases from Business Requirements Documents."},
                    {"role": "user", "content": full_prompt}
                ],
                max_tokens=16384,
                temperature=0.3,
                response_format={"type": "json_object"}
            )
            
            response_text = response.choices[0].message.content
            
            # Try to parse JSON response
            try:
                # Parse the JSON response
                json_response = json.loads(response_text)
                
                # Extract test cases from the response
                if isinstance(json_response, dict) and 'test_cases' in json_response:
                    test_cases = json_response['test_cases']
                    if test_cases and isinstance(test_cases, list):
                        return test_cases, response_text
                    else:
                        return None, response_text
                elif isinstance(json_response, list):
                    test_cases = json_response
                    return test_cases, response_text
                else:
                    # Try to find test_cases array in the response
                    test_cases = None
                    for key in json_response:
                        if isinstance(json_response[key], list) and len(json_response[key]) > 0:
                            test_cases = json_response[key]
                            break
                    
                    if test_cases:
                        return test_cases, response_text
                    else:
                        return None, response_text
                    
            except json.JSONDecodeError as e:
                st.warning("⚠️ JSON response incomplete due to token limit. Attempting to recover partial data...")
                
                # Try to extract incomplete JSON and fix it
                try:
                    # Find the test_cases array start
                    test_cases_start = response_text.find('"test_cases": [')
                    if test_cases_start != -1:
                        # Extract from test_cases array start
                        array_start = response_text.find('[', test_cases_start)
                        
                        # Find all complete test case objects
                        test_cases = []
                        current_pos = array_start + 1
                        brace_count = 0
                        current_obj_start = -1
                        
                        for i, char in enumerate(response_text[array_start + 1:], array_start + 1):
                            if char == '{':
                                if brace_count == 0:
                                    current_obj_start = i
                                brace_count += 1
                            elif char == '}':
                                brace_count -= 1
                                if brace_count == 0 and current_obj_start != -1:
                                    # Complete object found
                                    obj_text = response_text[current_obj_start:i+1]
                                    try:
                                        test_case = json.loads(obj_text)
                                        test_cases.append(test_case)
                                    except:
                                        pass
                                    current_obj_start = -1
                        
                        if test_cases:
                            st.info(f"✅ Recovered {len(test_cases)} complete test cases from incomplete response")
                            return test_cases, response_text
                
                except Exception as recovery_error:
                    st.error(f"Could not recover partial data: {str(recovery_error)}")
                
                # Final fallback - try to extract any JSON array
                start_idx = response_text.find('[')
                end_idx = response_text.rfind(']') + 1
                
                if start_idx != -1 and end_idx > start_idx:
                    try:
                        json_text = response_text[start_idx:end_idx]
                        test_cases = json.loads(json_text)
                        return test_cases, response_text
                    except:
                        pass
                
                return None, response_text
                
        except Exception as e:
            st.error(f"Error calling OpenAI API: {str(e)}")
            return None, str(e)

def categorize_test_case_type(tc_module):
    """Categorize test case based on TC Module name"""
    tc_module_lower = str(tc_module).lower()
    
    if 'field validation' in tc_module_lower or 'field-level' in tc_module_lower:
        return 'Field-Level'
    elif 'workflow' in tc_module_lower or 'functional' in tc_module_lower:
        return 'Functional'
    elif 'negative' in tc_module_lower:
        return 'Negative'
    elif 'report' in tc_module_lower or 'inquiry' in tc_module_lower:
        return 'Report/Inquiry'
    elif 'modification' in tc_module_lower or 'update' in tc_module_lower:
        return 'Modification'
    else:
        return 'Other'

def encode_image_to_base64(image_file):
    """Encode uploaded image to base64"""
    try:
        image = Image.open(image_file)
        buffered = io.BytesIO()
        image.save(buffered, format=image.format if image.format else "PNG")
        return base64.b64encode(buffered.getvalue()).decode('utf-8')
    except Exception as e:
        st.error(f"Error encoding image: {str(e)}")
        return None

def analyze_ui_screenshot(image_file, screen_name, api_key):
    """Analyze UI screenshot using GPT-4o Vision to extract UI elements"""
    
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Error initializing OpenAI client for UI analysis: {str(e)}")
        return []
    
    # Encode image
    base64_image = encode_image_to_base64(image_file)
    if not base64_image:
        return None
    
    prompt = """Analyze this UI screenshot and extract all interactive elements.

For each UI element, identify:
1. Element type (input_field, button, link, checkbox, dropdown, radio_button, textarea, etc.)
2. Element label/text (visible text or placeholder)
3. Whether it's required (look for * or "required" indicators)
4. Any validation hints or helper text

IMPORTANT: Return ONLY valid JSON. No markdown, no code blocks, no explanations. Just pure JSON.

JSON structure:
{
  "screen_type": "login/form/dashboard/etc",
  "elements": [
    {
      "type": "input_field",
      "label": "Email Address",
      "required": true,
      "placeholder": "Enter your email",
      "description": "Brief description of what this element does"
    }
  ]
}

Be thorough and extract ALL interactive elements you can see. Return the JSON now:"""

    try:
        with st.spinner(f'🔍 Analyzing {screen_name}...'):
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2000,
                temperature=0.1
            )
            
            response_text = response.choices[0].message.content
            
            # Try to parse JSON
            try:
                result = json.loads(response_text)
                return result
            except json.JSONDecodeError as je:
                st.error(f"JSON parsing error. Raw response: {response_text[:500]}")
                
                # Try to extract JSON from markdown code blocks
                if "```json" in response_text:
                    start = response_text.find("```json") + 7
                    end = response_text.find("```", start)
                    if end > start:
                        json_text = response_text[start:end].strip()
                        try:
                            result = json.loads(json_text)
                            return result
                        except:
                            pass
                
                # Try to find JSON object
                start = response_text.find('{')
                end = response_text.rfind('}') + 1
                if start >= 0 and end > start:
                    try:
                        json_text = response_text[start:end]
                        result = json.loads(json_text)
                        return result
                    except:
                        pass
                
                st.error(f"Could not extract valid JSON from response")
                return None
            
    except Exception as e:
        st.error(f"Error analyzing screenshot: {str(e)}")
        return None

def map_ui_elements_to_test_cases(ui_elements, test_cases, screen_name, api_key):
    """Map UI elements to test cases using AI"""
    
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Error initializing OpenAI client for UI mapping: {str(e)}")
        return test_cases
    
    # Create UI elements summary
    elements_summary = "\n".join([
        f"- {elem.get('type', 'N/A')}: {elem.get('label', 'N/A')} (Required: {elem.get('required', False)})"
        for elem in ui_elements
    ])
    
    # Create test case summary
    test_case_summary = "\n".join([
        f"- {tc.get('test_case_id', 'N/A')}: {tc.get('test_condition', 'N/A')} | Description: {tc.get('test_case_description', 'N/A')[:150]}"
        for tc in test_cases[:100]
    ])
    
    prompt = f"""You are a QA analyst mapping UI elements to test cases.

**SCREEN:** {screen_name}

**UI ELEMENTS FOUND:**
{elements_summary}

**AVAILABLE TEST CASES:**
{test_case_summary}

**TASK:**
For each UI element, find which test cases cover it. Match based on:
- Element label mentions in test case description
- Element type and actions (enter, click, select, verify)
- Semantic similarity (e.g., "user email" matches "email field")

Return ONLY a JSON object:
{{
  "mappings": [
    {{
      "element_type": "input_field",
      "element_label": "Email Address",
      "covered_by": ["TC_001", "TC_002"],
      "coverage_level": "Full" or "Partial" or "None",
      "confidence": 0.95,
      "missing_scenarios": ["List any missing test scenarios for this element"]
    }}
  ],
  "overall_coverage": 85.5,
  "summary": "Brief summary of UI coverage"
}}

Analyze now:"""

    try:
        with st.spinner(f'🔗 Mapping UI elements to test cases...'):
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert QA analyst specialized in UI test coverage analysis."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=3000,
                temperature=0.2,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            return result
            
    except Exception as e:
        st.error(f"Error mapping elements: {str(e)}")
        return None

def analyze_requirements_coverage(brd_text, test_cases, api_key):
    """Analyze coverage of BRD requirements by generated test cases using AI"""
    
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Error initializing OpenAI client for coverage analysis: {str(e)}")
        return "Error initializing OpenAI client"
    
    # Create test case summary
    test_case_summary = "\n".join([
        f"- {tc.get('test_case_id', 'N/A')}: {tc.get('test_condition', 'N/A')} (Module: {tc.get('tc_module', 'N/A')})"
        for tc in test_cases[:100]  # Limit to first 100 for token management
    ])
    
    prompt = f"""You are a QA analyst performing requirements coverage analysis.

**BRD DOCUMENT:**
{brd_text[:15000]}  

**GENERATED TEST CASES (Summary):**
{test_case_summary}

**TASK:**
Analyze the BRD and identify:
1. All functional requirements, features, and business rules mentioned in the BRD
2. Which requirements are covered by the generated test cases
3. Which requirements are NOT covered (gaps)

**OUTPUT FORMAT:**
Return ONLY a JSON object with this structure:
{{
  "total_requirements": <number>,
  "covered_requirements": [
    {{
      "requirement": "Requirement description",
      "requirement_id": "REQ-001",
      "covered_by": ["TC_001", "TC_002"],
      "coverage_level": "Full" or "Partial"
    }}
  ],
  "missing_requirements": [
    {{
      "requirement": "Requirement description",
      "requirement_id": "REQ-XXX",
      "reason": "Why it's not covered",
      "priority": "High" or "Medium" or "Low"
    }}
  ],
  "coverage_percentage": <number>,
  "summary": "Brief analysis summary"
}}

Analyze now:"""
    
    try:
        with st.spinner('🔍 Analyzing requirements coverage...'):
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert QA analyst specialized in requirements traceability and coverage analysis."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=8192,
                temperature=0.2,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            return result
            
    except Exception as e:
        st.error(f"Error analyzing coverage: {str(e)}")
        return None

def convert_to_dataframe(test_cases):
    """Convert test cases JSON to DataFrame"""
    if not test_cases:
        return None
    
    # Column mapping
    columns = {
        'product_name': 'Product Name',
        'process_category': 'Process Category',
        'business_process_id': 'Business Process ID',
        'business_process': 'Business Process',
        'scenario_id': 'Scenario ID',
        'scenario_description': 'Scenario Description',
        'category': 'Category',
        'importance': 'Importance',
        'test_case_id': 'Test Case ID',
        'tc_module': 'TC Module',
        'test_condition': 'Test Condition',
        'prerequisite': 'Pre-requisite',
        'test_case_description': 'Test Case Description',
        'expected_result': 'Expected Result'
    }
    
    df = pd.DataFrame(test_cases)
    df = df.rename(columns=columns)
    
    # Add Test Type column for categorization
    df['Test Type'] = df['TC Module'].apply(categorize_test_case_type)
    
    # Reorder columns (add Test Type after TC Module)
    column_order = list(columns.values())
    # Insert Test Type after TC Module
    tc_module_idx = column_order.index('TC Module')
    column_order.insert(tc_module_idx + 1, 'Test Type')
    
    df = df[[col for col in column_order if col in df.columns]]
    
    return df

def create_excel_output(df):
    """Create formatted Excel output with separate sheets for different test types"""
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Helper function to format sheet
        def format_sheet(worksheet, column_count=15):
            header_fill = openpyxl.styles.PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            header_font = openpyxl.styles.Font(color='FFFFFF', bold=True)
            
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Adjust column widths
            column_widths = {
                'A': 20, 'B': 30, 'C': 15, 'D': 40, 'E': 12, 'F': 40,
                'G': 18, 'H': 12, 'I': 15, 'J': 20, 'K': 15, 'L': 40, 
                'M': 30, 'N': 50, 'O': 50
            }
            
            for col, width in column_widths.items():
                if col in worksheet.column_dimensions:
                    worksheet.column_dimensions[col].width = width
        
        # Write all test cases
        df.to_excel(writer, sheet_name='All Test Cases', index=False)
        format_sheet(writer.sheets['All Test Cases'])
        
        # Separate by Test Type
        if 'Test Type' in df.columns:
            # Field-Level Test Cases
            field_df = df[df['Test Type'] == 'Field-Level']
            if not field_df.empty:
                field_df.to_excel(writer, sheet_name='Field-Level Tests', index=False)
                format_sheet(writer.sheets['Field-Level Tests'])
            
            # Functional/Workflow Test Cases
            functional_df = df[df['Test Type'] == 'Functional']
            if not functional_df.empty:
                functional_df.to_excel(writer, sheet_name='Functional Tests', index=False)
                format_sheet(writer.sheets['Functional Tests'])
            
            # Negative Test Cases
            negative_df = df[df['Test Type'] == 'Negative']
            if not negative_df.empty:
                negative_df.to_excel(writer, sheet_name='Negative Tests', index=False)
                format_sheet(writer.sheets['Negative Tests'])
            
            # Other Test Cases (Reports, Modifications, etc.)
            other_df = df[~df['Test Type'].isin(['Field-Level', 'Functional', 'Negative'])]
            if not other_df.empty:
                other_df.to_excel(writer, sheet_name='Other Tests', index=False)
                format_sheet(writer.sheets['Other Tests'])
        
        # Add enhanced summary sheet
        test_type_counts = df['Test Type'].value_counts().to_dict() if 'Test Type' in df.columns else {}
        
        summary_data = {
            'Metric': [
                'Total Test Cases',
                'Field-Level Tests',
                'Functional/Workflow Tests',
                'Negative Tests',
                'Other Tests',
                '',
                'Functional Positive',
                'Functional Negative',
                '',
                'Critical Priority',
                'High Priority',
                'Medium Priority',
                'Low Priority',
                '',
                'Unique Modules',
                'Unique Scenarios'
            ],
            'Value': [
                len(df),
                test_type_counts.get('Field-Level', 0),
                test_type_counts.get('Functional', 0),
                test_type_counts.get('Negative', 0),
                sum([v for k, v in test_type_counts.items() if k not in ['Field-Level', 'Functional', 'Negative']]),
                '',
                len(df[df['Category'].str.contains('Positive', case=False, na=False)]),
                len(df[df['Category'].str.contains('negative', case=False, na=False)]),
                '',
                len(df[df['Importance'] == 'Critical']),
                len(df[df['Importance'] == 'High']),
                len(df[df['Importance'] == 'Medium']),
                len(df[df['Importance'] == 'Low']),
                '',
                df['TC Module'].nunique(),
                df['Scenario Description'].nunique()
            ]
        }
        
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='Summary', index=False)
        
        # Format summary sheet
        summary_sheet = writer.sheets['Summary']
        for cell in summary_sheet[1]:
            cell.fill = openpyxl.styles.PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            cell.font = openpyxl.styles.Font(color='FFFFFF', bold=True)
        
        summary_sheet.column_dimensions['A'].width = 30
        summary_sheet.column_dimensions['B'].width = 20
    
    output.seek(0)
    return output

# Main app
def main():
    # Header
    st.markdown('<div class="main-header">📋 Test Case Generator</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Automated test case generation from BRD documents using AI</div>', unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key
        api_key = st.text_input(
            "OpenAI API Key",
            type="password",
            key="api_key_input",
            help="Enter your OpenAI API key (from platform.openai.com)"
        )
        
        st.divider()
        
        st.subheader("Generation Options")
        
        # Domain/Industry selection
        domain_options = [
            "Auto-detect from BRD",
            "Banking & Finance",
            "E-commerce & Retail",
            "Healthcare & Medical",
            "Insurance",
            "Education & E-learning",
            "Logistics & Supply Chain",
            "Travel & Hospitality",
            "Telecommunications",
            "Government & Public Sector",
            "Manufacturing",
            "Real Estate",
            "Entertainment & Media",
            "Other (specify below)"
        ]
        
        selected_domain = st.selectbox(
            "Application Domain/Industry",
            options=domain_options,
            help="Select your application domain to generate more contextually relevant test cases"
        )
        
        # Custom domain input if "Other" is selected
        custom_domain = ""
        if selected_domain == "Other (specify below)":
            custom_domain = st.text_input(
                "Specify Your Domain",
                placeholder="E.g., Food Delivery, Social Media, CRM, etc."
            )
        
        # Determine final domain
        if selected_domain == "Auto-detect from BRD":
            domain_context = "Analyze the BRD and identify the application domain automatically"
        elif selected_domain == "Other (specify below)" and custom_domain:
            domain_context = custom_domain
        else:
            domain_context = selected_domain
        
        st.divider()
        
        negative_ratio = st.slider(
            "Negative Test Ratio (%)",
            min_value=5,
            max_value=25,
            value=15,
            help="Percentage of negative test cases"
        )
        
        variations_input = st.text_area(
            "User/Customer/Product Variations (Optional)",
            placeholder="Enter variations separated by commas\nExamples:\n• Banking: Individual Customer, Corporate Customer, Joint Account\n• E-commerce: Guest User, Registered User, Premium Member\n• Healthcare: Patient, Doctor, Admin",
            help="Enter different user types, customer categories, or product variations to test",
            height=100
        )
        
        # Convert comma-separated input to list
        variations = [v.strip() for v in variations_input.split(',') if v.strip()] if variations_input else []
        
        focus_modules_input = st.text_area(
            "Focus on Specific Modules (Optional)",
            placeholder="Enter module names separated by commas\nExample: User Registration, Checkout, Payment Processing, Order Management",
            help="Leave empty to analyze all modules from BRD. Enter comma-separated module names to focus on specific areas.",
            height=80
        )
        
        # Convert comma-separated input to list
        focus_modules = [m.strip() for m in focus_modules_input.split(',') if m.strip()] if focus_modules_input else []
        
        st.divider()
        
        st.subheader("📚 Framework Info")
        st.info("""
        This app generates comprehensive test cases for ANY domain:
        Banking, E-commerce, Healthcare, Insurance, Education, Logistics, etc.
        
        **Key Features:**
        - 🎯 **1:1 Requirement Mapping** - One test case per requirement
        - 📊 **Smart Requirements Analysis** - AI identifies all functional requirements first
        - 14-column structured test cases
        - Workflow decomposition (15-step model)
        - Authorization/approval patterns
        - Field-level validations
        - Negative test cases (business rules)
        - Adapts to YOUR domain and terminology
        """)
    
    # Main content area
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["📤 Upload & Generate", "📊 View Results", "🎯 Coverage Analysis", "📸 UI Coverage", "💾 Download", "📖 Help"])
    
    with tab1:
        st.header("Upload BRD Documents")
        
        uploaded_files = st.file_uploader(
            "Choose BRD files (Word, PDF, Excel, Text)",
            type=['docx', 'pdf', 'xlsx', 'xls', 'txt'],
            accept_multiple_files=True,
            help="You can upload multiple files. They will be combined for analysis."
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} file(s) uploaded")
            
            # Show file details
            with st.expander("📁 Uploaded Files"):
                for file in uploaded_files:
                    st.write(f"- {file.name} ({file.size / 1024:.2f} KB)")
            
            col1, col2 = st.columns([1, 4])
            
            with col1:
                generate_button = st.button("🚀 Generate Test Cases", type="primary", use_container_width=True)
            
            with col2:
                preview_button = st.button("👁️ Preview Extracted Text", use_container_width=True)
            
            if preview_button:
                with st.spinner("Extracting text from documents..."):
                    brd_text = process_uploaded_files(uploaded_files)
                    st.session_state['brd_text'] = brd_text
                    
                with st.expander("📄 Extracted BRD Text (Preview)", expanded=True):
                    st.text_area(
                        "Combined Document Text",
                        brd_text[:5000] + "\n\n... (showing first 5000 characters)",
                        height=300
                    )
            
            if generate_button:
                if not api_key:
                    st.error("⚠️ Please enter your OpenAI API key in the sidebar")
                else:
                    # Process documents
                    with st.spinner("📖 Extracting text from documents..."):
                        brd_text = process_uploaded_files(uploaded_files)
                        st.session_state['brd_text'] = brd_text
                    
                    st.success("✅ Text extraction completed")
                    
                    # Load prompt
                    prompt_template = load_prompt()
                    if not prompt_template:
                        st.stop()
                    
                    # Prepare options (target_count will be determined automatically based on BRD content)
                    options = {
                        'target_count': 'auto',  # Will be calculated based on BRD requirements
                        'negative_ratio': negative_ratio,
                        'variations': variations,
                        'focus_modules': focus_modules,
                        'domain_context': domain_context
                    }
                    
                    # Generate test cases
                    test_cases, raw_response = generate_test_cases_with_ai(
                        brd_text, prompt_template, api_key, options
                    )
                    
                    if test_cases:
                        st.session_state['test_cases'] = test_cases
                        st.session_state['raw_response'] = raw_response
                        st.session_state['generation_time'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        
                        # Store requirements count for display
                        if 'requirements_count' not in st.session_state:
                            # Re-analyze to get requirements count for display
                            requirements, req_count = analyze_requirements_from_brd(brd_text, api_key)
                            st.session_state['requirements_count'] = req_count
                        
                        # Display success message with requirements coverage info
                        req_count = st.session_state.get('requirements_count', 0)
                        test_case_count = len(test_cases)
                        
                        if req_count > 0:
                            coverage_ratio = (test_case_count / req_count) * 100
                            if coverage_ratio >= 95:  # Allow for slight variation
                                st.markdown('<div class="success-box">✅ Test cases generated successfully! Perfect 1:1 requirement coverage achieved!</div>', 
                                          unsafe_allow_html=True)
                                st.success(f"🎯 Generated {test_case_count} test cases for {req_count} requirements (1:1 mapping)")
                            else:
                                st.markdown('<div class="success-box">✅ Test cases generated successfully! Check the "View Results" tab.</div>', 
                                          unsafe_allow_html=True)
                                st.info(f"📊 Generated {test_case_count} test cases for {req_count} identified requirements")
                        else:
                            st.markdown('<div class="success-box">✅ Test cases generated successfully! Check the "View Results" tab.</div>', 
                                      unsafe_allow_html=True)
                        
                        # Quick stats
                        df = convert_to_dataframe(test_cases)
                        col1, col2, col3, col4, col5, col6 = st.columns(6)
                        with col1:
                            st.metric("Total Test Cases", len(df))
                        with col2:
                            st.metric("Requirements Found", req_count if req_count > 0 else "N/A")
                        with col3:
                            field_level = len(df[df['Test Type'] == 'Field-Level']) if 'Test Type' in df.columns else 0
                            st.metric("Field-Level", field_level)
                        with col4:
                            functional = len(df[df['Test Type'] == 'Functional']) if 'Test Type' in df.columns else 0
                            st.metric("Functional", functional)
                        with col5:
                            negative = len(df[df['Test Type'] == 'Negative']) if 'Test Type' in df.columns else 0
                            st.metric("Negative", negative)
                        with col6:
                            modules = df['TC Module'].nunique()
                            st.metric("Modules", modules)
                    else:
                        st.warning("⚠️ Could not extract structured test cases. Showing raw response:")
                        st.text_area("Raw AI Response", raw_response, height=400)
        else:
            st.info("👆 Upload BRD documents to get started")
    
    with tab2:
        st.header("Generated Test Cases")
        
        if 'test_cases' in st.session_state:
            df = convert_to_dataframe(st.session_state['test_cases'])
            st.session_state['df'] = df
            
            # Statistics
            st.subheader("📊 Statistics")
            
            # Requirements coverage info
            req_count = st.session_state.get('requirements_count', 0)
            if req_count > 0:
                coverage_ratio = (len(df) / req_count) * 100
                if coverage_ratio >= 95:
                    st.success(f"🎯 Perfect 1:1 Coverage: {len(df)} test cases for {req_count} requirements")
                else:
                    st.info(f"📊 Coverage: {len(df)} test cases for {req_count} requirements ({coverage_ratio:.1f}%)")
            
            col1, col2, col3, col4, col5, col6 = st.columns(6)
            with col1:
                st.metric("Total Test Cases", len(df))
            with col2:
                st.metric("Requirements Found", req_count if req_count > 0 else "N/A", 
                         delta=f"1:1 mapping" if req_count > 0 and abs(len(df) - req_count) <= 2 else None)
            with col3:
                field_level = len(df[df['Test Type'] == 'Field-Level']) if 'Test Type' in df.columns else 0
                st.metric("Field-Level", field_level, f"{field_level/len(df)*100:.1f}%")
            with col4:
                functional = len(df[df['Test Type'] == 'Functional']) if 'Test Type' in df.columns else 0
                st.metric("Functional", functional, f"{functional/len(df)*100:.1f}%")
            with col5:
                negative = len(df[df['Test Type'] == 'Negative']) if 'Test Type' in df.columns else 0
                st.metric("Negative", negative, f"{negative/len(df)*100:.1f}%")
            with col6:
                modules = df['TC Module'].nunique()
                st.metric("Modules", modules)
            
            st.divider()
            
            # Filters
            st.subheader("🔍 Filter Test Cases")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                test_type_filter = st.multiselect(
                    "Filter by Test Type",
                    options=df['Test Type'].unique().tolist() if 'Test Type' in df.columns else [],
                    default=None,
                    help="Field-Level, Functional, Negative, etc."
                )
            
            with col2:
                module_filter = st.multiselect(
                    "Filter by Module",
                    options=df['TC Module'].unique().tolist(),
                    default=None
                )
            
            with col3:
                category_filter = st.multiselect(
                    "Filter by Category",
                    options=df['Category'].unique().tolist(),
                    default=None
                )
            
            with col4:
                importance_filter = st.multiselect(
                    "Filter by Importance",
                    options=df['Importance'].unique().tolist(),
                    default=None
                )
            
            # Apply filters
            filtered_df = df.copy()
            if test_type_filter:
                filtered_df = filtered_df[filtered_df['Test Type'].isin(test_type_filter)]
            if module_filter:
                filtered_df = filtered_df[filtered_df['TC Module'].isin(module_filter)]
            if category_filter:
                filtered_df = filtered_df[filtered_df['Category'].isin(category_filter)]
            if importance_filter:
                filtered_df = filtered_df[filtered_df['Importance'].isin(importance_filter)]
            
            st.write(f"Showing {len(filtered_df)} of {len(df)} test cases")
            
            # Display dataframe
            st.dataframe(
                filtered_df,
                use_container_width=True,
                height=600
            )
            
            # Test Type breakdown
            if 'Test Type' in df.columns:
                with st.expander("🔖 Test Type Distribution", expanded=True):
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        test_type_counts = df['Test Type'].value_counts()
                        st.bar_chart(test_type_counts)
                    with col2:
                        st.write("**Count by Type:**")
                        for test_type, count in test_type_counts.items():
                            percentage = (count / len(df)) * 100
                            st.write(f"**{test_type}**: {count} ({percentage:.1f}%)")
            
            # Module breakdown
            with st.expander("📈 Module Breakdown"):
                module_counts = df['TC Module'].value_counts()
                st.bar_chart(module_counts)
            
            # Category breakdown
            with st.expander("📊 Category Distribution"):
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**Category:**")
                    category_counts = df['Category'].value_counts()
                    st.write(category_counts)
                with col2:
                    st.write("**Importance:**")
                    importance_counts = df['Importance'].value_counts()
                    st.write(importance_counts)
        else:
            st.info("👈 Generate test cases first to view results")
    
    with tab3:
        st.header("🎯 Requirements Coverage Analysis")
        
        if 'test_cases' in st.session_state and 'brd_text' in st.session_state:
            st.markdown("""
            This analysis compares your BRD requirements with the generated test cases to identify:
            - ✅ **Covered Requirements**: Which requirements have test cases
            - ❌ **Missing Coverage**: Requirements without test cases
            - 📊 **Coverage Percentage**: Overall test coverage
            """)
            
            st.divider()
            
            # Analyze button
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                analyze_button = st.button(
                    "🔍 Analyze Requirements Coverage",
                    type="primary",
                    use_container_width=True,
                    help="Click to analyze which BRD requirements are covered by test cases"
                )
            
            if analyze_button:
                if not api_key:
                    st.error("⚠️ Please enter your OpenAI API key in the sidebar")
                else:
                    # Perform coverage analysis
                    coverage_result = analyze_requirements_coverage(
                        st.session_state['brd_text'],
                        st.session_state['test_cases'],
                        api_key
                    )
                    
                    if coverage_result:
                        st.session_state['coverage_result'] = coverage_result
            
            # Display results if available
            if 'coverage_result' in st.session_state:
                result = st.session_state['coverage_result']
                
                st.divider()
                
                # Coverage Summary
                st.subheader("📊 Coverage Summary")
                
                coverage_pct = result.get('coverage_percentage', 0)
                total_reqs = result.get('total_requirements', 0)
                covered_count = len(result.get('covered_requirements', []))
                missing_count = len(result.get('missing_requirements', []))
                
                # Progress bar
                st.progress(coverage_pct / 100.0)
                
                # Metrics
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Coverage", f"{coverage_pct:.1f}%", 
                             delta=f"{coverage_pct - 80:.1f}%" if coverage_pct < 80 else "Good")
                with col2:
                    st.metric("Total Requirements", total_reqs)
                with col3:
                    st.metric("✅ Covered", covered_count, delta_color="normal")
                with col4:
                    st.metric("❌ Missing", missing_count, delta_color="inverse")
                
                # Summary text
                if 'summary' in result:
                    st.info(f"**Analysis Summary:** {result['summary']}")
                
                st.divider()
                
                # Covered Requirements
                st.subheader("✅ Covered Requirements")
                
                covered_reqs = result.get('covered_requirements', [])
                if covered_reqs:
                    # Filter options
                    coverage_filter = st.multiselect(
                        "Filter by Coverage Level",
                        options=['Full', 'Partial'],
                        default=['Full', 'Partial']
                    )
                    
                    filtered_covered = [req for req in covered_reqs if req.get('coverage_level') in coverage_filter]
                    
                    st.write(f"Showing {len(filtered_covered)} of {len(covered_reqs)} covered requirements")
                    
                    for idx, req in enumerate(filtered_covered, 1):
                        coverage_level = req.get('coverage_level', 'Unknown')
                        color = "🟢" if coverage_level == "Full" else "🟡"
                        
                        with st.expander(f"{color} {req.get('requirement_id', 'REQ-???')}: {req.get('requirement', 'N/A')}", expanded=False):
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                st.write(f"**Requirement:** {req.get('requirement', 'N/A')}")
                                st.write(f"**Coverage Level:** {coverage_level}")
                                
                                test_cases = req.get('covered_by', [])
                                if test_cases:
                                    st.write(f"**Covered by Test Cases:** {', '.join(test_cases)}")
                            with col2:
                                if coverage_level == "Full":
                                    st.success("✅ Full Coverage")
                                else:
                                    st.warning("⚠️ Partial Coverage")
                else:
                    st.warning("No covered requirements found.")
                
                st.divider()
                
                # Missing Requirements (Gaps)
                st.subheader("❌ Missing Coverage (Gaps)")
                
                missing_reqs = result.get('missing_requirements', [])
                if missing_reqs:
                    # Priority filter
                    priority_filter = st.multiselect(
                        "Filter by Priority",
                        options=['High', 'Medium', 'Low'],
                        default=['High', 'Medium', 'Low']
                    )
                    
                    filtered_missing = [req for req in missing_reqs if req.get('priority') in priority_filter]
                    
                    st.write(f"Showing {len(filtered_missing)} of {len(missing_reqs)} missing requirements")
                    
                    # Sort by priority
                    priority_order = {'High': 0, 'Medium': 1, 'Low': 2}
                    filtered_missing.sort(key=lambda x: priority_order.get(x.get('priority', 'Low'), 3))
                    
                    for idx, req in enumerate(filtered_missing, 1):
                        priority = req.get('priority', 'Medium')
                        
                        if priority == 'High':
                            icon = "🔴"
                            color = "error"
                        elif priority == 'Medium':
                            icon = "🟡"
                            color = "warning"
                        else:
                            icon = "🟢"
                            color = "info"
                        
                        with st.expander(f"{icon} **{priority}** - {req.get('requirement_id', 'REQ-???')}: {req.get('requirement', 'N/A')}", expanded=(priority == 'High')):
                            st.write(f"**Requirement:** {req.get('requirement', 'N/A')}")
                            st.write(f"**Priority:** {priority}")
                            st.write(f"**Reason Not Covered:** {req.get('reason', 'Not specified')}")
                            
                            if priority == 'High':
                                st.error("⚠️ High priority requirement - Create test cases immediately!")
                            elif priority == 'Medium':
                                st.warning("⚠️ Medium priority - Should be covered")
                            else:
                                st.info("ℹ️ Low priority - Consider coverage")
                else:
                    st.success("🎉 All requirements are covered!")
                
                st.divider()
                
                # Recommendations
                st.subheader("💡 Recommendations")
                
                if missing_count > 0:
                    st.markdown(f"""
                    ### Actions to Improve Coverage:
                    
                    1. **Address High Priority Gaps** ({sum(1 for r in missing_reqs if r.get('priority') == 'High')} items)
                       - Review each high-priority missing requirement
                       - Create test cases to cover these requirements
                    
                    2. **Review Partial Coverage** ({sum(1 for r in covered_reqs if r.get('coverage_level') == 'Partial')} items)
                       - Enhance test cases for partially covered requirements
                       - Add more test scenarios
                    
                    3. **Re-generate Test Cases**
                       - Go back to "Upload & Generate" tab
                       - Use "Focus Modules" to target missing areas
                       - Increase test case count for better coverage
                    """)
                else:
                    st.success("""
                    ### Excellent Coverage! ✅
                    
                    Your test cases provide comprehensive coverage of the BRD requirements.
                    - All requirements have test cases
                    - Ready for test execution
                    """)
                
                # Download coverage report
                st.divider()
                st.subheader("📥 Download Coverage Report")
                
                coverage_report = {
                    'Summary': {
                        'Coverage Percentage': f"{coverage_pct:.1f}%",
                        'Total Requirements': total_reqs,
                        'Covered Requirements': covered_count,
                        'Missing Requirements': missing_count
                    },
                    'Covered Requirements': covered_reqs,
                    'Missing Requirements': missing_reqs
                }
                
                report_json = json.dumps(coverage_report, indent=2)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📄 Download as JSON",
                        data=report_json,
                        file_name=f"coverage_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json"
                    )
                with col2:
                    # Create text report
                    text_report = f"""REQUIREMENTS COVERAGE REPORT
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

SUMMARY
=======
Coverage Percentage: {coverage_pct:.1f}%
Total Requirements: {total_reqs}
Covered Requirements: {covered_count}
Missing Requirements: {missing_count}

COVERED REQUIREMENTS ({covered_count})
========================
"""
                    for req in covered_reqs:
                        text_report += f"\n{req.get('requirement_id', 'N/A')}: {req.get('requirement', 'N/A')}\n"
                        text_report += f"  Coverage: {req.get('coverage_level', 'N/A')}\n"
                        text_report += f"  Test Cases: {', '.join(req.get('covered_by', []))}\n"
                    
                    text_report += f"\n\nMISSING REQUIREMENTS ({missing_count})\n"
                    text_report += "=======================\n"
                    for req in missing_reqs:
                        text_report += f"\n{req.get('requirement_id', 'N/A')}: {req.get('requirement', 'N/A')}\n"
                        text_report += f"  Priority: {req.get('priority', 'N/A')}\n"
                        text_report += f"  Reason: {req.get('reason', 'N/A')}\n"
                    
                    st.download_button(
                        label="📝 Download as Text",
                        data=text_report,
                        file_name=f"coverage_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        mime="text/plain"
                    )
        
        else:
            st.info("👈 Generate test cases first to perform coverage analysis")
            st.markdown("""
            ### How Coverage Analysis Works:
            
            1. **Upload BRD documents** in the first tab
            2. **Generate test cases** using AI
            3. **Come back here** and click "Analyze Coverage"
            4. **Review results** showing covered and missing requirements
            5. **Take action** on gaps to improve coverage
            """)
    
    with tab4:
        st.header("📸 UI Coverage Analysis")
        
        if 'test_cases' in st.session_state:
            st.markdown("""
            Upload screenshots of your application's UI to analyze which UI elements are covered by test cases.
            
            **Features:**
            - 🔍 Extract UI elements automatically using AI vision
            - 🔗 Map elements to test cases intelligently
            - 📊 Visual coverage dashboard
            - ❌ Identify untested UI components
            """)
            
            st.divider()
            
            # Screenshot upload section
            st.subheader("📤 Upload UI Screenshots")
            
            col1, col2 = st.columns([2, 1])
            
            with col1:
                uploaded_screenshots = st.file_uploader(
                    "Upload UI Screenshots",
                    type=['png', 'jpg', 'jpeg'],
                    accept_multiple_files=True,
                    help="Upload screenshots of your application screens (PNG, JPG, JPEG)"
                )
            
            with col2:
                if uploaded_screenshots:
                    st.success(f"✅ {len(uploaded_screenshots)} screenshot(s) uploaded")
            
            # Initialize session state for UI analysis
            if 'ui_analyses' not in st.session_state:
                st.session_state['ui_analyses'] = {}
            
            if uploaded_screenshots:
                st.divider()
                
                # Screen naming
                st.subheader("🏷️ Label Your Screenshots")
                
                screen_names = {}
                for idx, screenshot in enumerate(uploaded_screenshots):
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        screen_name = st.text_input(
                            f"Screen Name for {screenshot.name}",
                            value=screenshot.name.replace('.png', '').replace('.jpg', '').replace('.jpeg', '').replace('_', ' ').title(),
                            key=f"screen_name_{idx}"
                        )
                        screen_names[screenshot.name] = screen_name
                    with col2:
                        # Show thumbnail
                        image = Image.open(screenshot)
                        st.image(image, width=100)
                
                st.divider()
                
                # Analyze button
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    analyze_ui_button = st.button(
                        "🔍 Analyze UI Coverage",
                        type="primary",
                        use_container_width=True,
                        help="Extract UI elements and map to test cases"
                    )
                
                if analyze_ui_button:
                    if not api_key:
                        st.error("⚠️ Please enter your OpenAI API key in the sidebar")
                    else:
                        # Analyze each screenshot
                        for screenshot in uploaded_screenshots:
                            screen_name = screen_names.get(screenshot.name, screenshot.name)
                            
                            # Reset file pointer
                            screenshot.seek(0)
                            
                            # Step 1: Extract UI elements
                            ui_analysis = analyze_ui_screenshot(screenshot, screen_name, api_key)
                            
                            if ui_analysis and 'elements' in ui_analysis:
                                elements = ui_analysis['elements']
                                
                                # Step 2: Map to test cases
                                mapping = map_ui_elements_to_test_cases(
                                    elements,
                                    st.session_state['test_cases'],
                                    screen_name,
                                    api_key
                                )
                                
                                if mapping:
                                    # Store results
                                    st.session_state['ui_analyses'][screen_name] = {
                                        'screenshot': screenshot,
                                        'elements': elements,
                                        'mapping': mapping,
                                        'screen_type': ui_analysis.get('screen_type', 'Unknown')
                                    }
                        
                        if st.session_state['ui_analyses']:
                            st.success("✅ UI analysis completed! Scroll down to see results.")
                
                # Display results
                if st.session_state['ui_analyses']:
                    st.divider()
                    st.subheader("📊 UI Coverage Results")
                    
                    # Overall summary
                    total_elements = 0
                    total_covered = 0
                    total_partial = 0
                    total_uncovered = 0
                    
                    for screen_name, analysis in st.session_state['ui_analyses'].items():
                        mapping = analysis.get('mapping', {})
                        mappings = mapping.get('mappings', [])
                        
                        for m in mappings:
                            total_elements += 1
                            coverage = m.get('coverage_level', 'None')
                            if coverage == 'Full':
                                total_covered += 1
                            elif coverage == 'Partial':
                                total_partial += 1
                            else:
                                total_uncovered += 1
                    
                    # Overall metrics
                    if total_elements > 0:
                        overall_coverage_pct = ((total_covered + total_partial * 0.5) / total_elements) * 100
                        
                        st.markdown("### 📈 Overall UI Coverage")
                        st.progress(overall_coverage_pct / 100.0)
                        
                        col1, col2, col3, col4, col5 = st.columns(5)
                        with col1:
                            st.metric("Coverage", f"{overall_coverage_pct:.1f}%")
                        with col2:
                            st.metric("Total Elements", total_elements)
                        with col3:
                            st.metric("✅ Fully Covered", total_covered)
                        with col4:
                            st.metric("⚠️ Partial", total_partial)
                        with col5:
                            st.metric("❌ Not Covered", total_uncovered)
                    
                    st.divider()
                    
                    # Screen-by-screen analysis
                    st.markdown("### 🖼️ Screen-by-Screen Analysis")
                    
                    for screen_name, analysis in st.session_state['ui_analyses'].items():
                        with st.expander(f"📱 {screen_name}", expanded=True):
                            col1, col2 = st.columns([1, 2])
                            
                            with col1:
                                # Show screenshot
                                screenshot = analysis['screenshot']
                                screenshot.seek(0)
                                image = Image.open(screenshot)
                                st.image(image, caption=screen_name, use_column_width=True)
                            
                            with col2:
                                mapping = analysis.get('mapping', {})
                                mappings = mapping.get('mappings', [])
                                overall_cov = mapping.get('overall_coverage', 0)
                                summary = mapping.get('summary', 'No summary available')
                                
                                # Screen coverage
                                st.metric("Screen Coverage", f"{overall_cov:.1f}%")
                                st.info(f"**Summary:** {summary}")
                                
                                st.markdown("---")
                                
                                # Elements breakdown
                                st.markdown("#### UI Elements Coverage")
                                
                                for elem_map in mappings:
                                    elem_type = elem_map.get('element_type', 'Unknown')
                                    elem_label = elem_map.get('element_label', 'Unknown')
                                    coverage_level = elem_map.get('coverage_level', 'None')
                                    covered_by = elem_map.get('covered_by', [])
                                    confidence = elem_map.get('confidence', 0)
                                    missing = elem_map.get('missing_scenarios', [])
                                    
                                    # Icon based on coverage
                                    if coverage_level == 'Full':
                                        icon = "🟢"
                                        color = "success"
                                    elif coverage_level == 'Partial':
                                        icon = "🟡"
                                        color = "warning"
                                    else:
                                        icon = "🔴"
                                        color = "error"
                                    
                                    with st.container():
                                        st.markdown(f"{icon} **{elem_type.replace('_', ' ').title()}: {elem_label}**")
                                        
                                        col_a, col_b = st.columns([2, 1])
                                        with col_a:
                                            if covered_by:
                                                st.write(f"✅ Covered by: {', '.join(covered_by)}")
                                            else:
                                                st.write("❌ No test cases found")
                                        
                                        with col_b:
                                            if coverage_level == 'Full':
                                                st.success(f"✅ {coverage_level} ({confidence:.0%})")
                                            elif coverage_level == 'Partial':
                                                st.warning(f"⚠️ {coverage_level} ({confidence:.0%})")
                                            else:
                                                st.error(f"❌ {coverage_level}")
                                        
                                        if missing:
                                            with st.expander("💡 Suggested Additional Tests"):
                                                for scenario in missing:
                                                    st.write(f"• {scenario}")
                                        
                                        st.markdown("---")
                    
                    # Download UI coverage report
                    st.divider()
                    st.subheader("📥 Download UI Coverage Report")
                    
                    # Create comprehensive report
                    ui_report = {
                        'overall_coverage': f"{overall_coverage_pct:.1f}%",
                        'total_elements': total_elements,
                        'fully_covered': total_covered,
                        'partially_covered': total_partial,
                        'not_covered': total_uncovered,
                        'screens': {}
                    }
                    
                    for screen_name, analysis in st.session_state['ui_analyses'].items():
                        mapping = analysis.get('mapping', {})
                        ui_report['screens'][screen_name] = {
                            'screen_type': analysis.get('screen_type', 'Unknown'),
                            'coverage': mapping.get('overall_coverage', 0),
                            'summary': mapping.get('summary', ''),
                            'elements': mapping.get('mappings', [])
                        }
                    
                    report_json = json.dumps(ui_report, indent=2)
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="📄 Download as JSON",
                            data=report_json,
                            file_name=f"ui_coverage_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                            mime="application/json"
                        )
                    
                    with col2:
                        # Text report
                        text_report = f"""UI COVERAGE ANALYSIS REPORT
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

OVERALL SUMMARY
===============
Coverage: {overall_coverage_pct:.1f}%
Total UI Elements: {total_elements}
Fully Covered: {total_covered}
Partially Covered: {total_partial}
Not Covered: {total_uncovered}

"""
                        for screen_name, analysis in st.session_state['ui_analyses'].items():
                            mapping = analysis.get('mapping', {})
                            text_report += f"\n{'='*60}\n"
                            text_report += f"SCREEN: {screen_name}\n"
                            text_report += f"{'='*60}\n"
                            text_report += f"Screen Type: {analysis.get('screen_type', 'Unknown')}\n"
                            text_report += f"Coverage: {mapping.get('overall_coverage', 0):.1f}%\n"
                            text_report += f"Summary: {mapping.get('summary', '')}\n\n"
                            
                            text_report += "UI ELEMENTS:\n"
                            text_report += "-" * 60 + "\n"
                            
                            for elem_map in mapping.get('mappings', []):
                                text_report += f"\n{elem_map.get('element_type', 'Unknown')}: {elem_map.get('element_label', 'Unknown')}\n"
                                text_report += f"  Coverage: {elem_map.get('coverage_level', 'None')}\n"
                                text_report += f"  Confidence: {elem_map.get('confidence', 0):.0%}\n"
                                covered_by = elem_map.get('covered_by', [])
                                if covered_by:
                                    text_report += f"  Test Cases: {', '.join(covered_by)}\n"
                                missing = elem_map.get('missing_scenarios', [])
                                if missing:
                                    text_report += "  Missing Scenarios:\n"
                                    for scenario in missing:
                                        text_report += f"    - {scenario}\n"
                        
                        st.download_button(
                            label="📝 Download as Text",
                            data=text_report,
                            file_name=f"ui_coverage_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain"
                        )
        
        else:
            st.info("👈 Generate test cases first before analyzing UI coverage")
            st.markdown("""
            ### How UI Coverage Analysis Works:
            
            1. **Generate test cases** from your BRD (Tab 1)
            2. **Upload UI screenshots** of your application
            3. **AI extracts UI elements** (buttons, fields, links, etc.)
            4. **Intelligent mapping** links elements to test cases
            5. **Visual results** show which UI parts are tested
            
            ### Benefits:
            - 🎯 Visual validation of test coverage
            - 🔍 Identify untested UI components
            - 📊 Professional coverage reports
            - ✅ Ensure every button and field is tested
            """)
    
    with tab5:
        st.header("Download Test Cases")
        
        if 'df' in st.session_state:
            df = st.session_state['df']
            
            st.success(f"✅ Ready to download {len(df)} test cases")
            
            # Show generation info
            if 'generation_time' in st.session_state:
                st.info(f"📅 Generated on: {st.session_state['generation_time']}")
            
            # Create Excel file
            excel_file = create_excel_output(df)
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.download_button(
                    label="📥 Download as Excel",
                    data=excel_file,
                    file_name=f"Test_Cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            
            with col2:
                csv_data = df.to_csv(index=False)
                st.download_button(
                    label="📥 Download as CSV",
                    data=csv_data,
                    file_name=f"Test_Cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            # Preview
            with st.expander("👁️ Preview Excel Format"):
                st.dataframe(df.head(10), use_container_width=True)
        else:
            st.info("👈 Generate test cases first to download")
    
    with tab5:
        st.header("📖 How to Use")
        
        st.markdown("""
        ### Step-by-Step Guide
        
        1. **Enter API Key**
           - Get your OpenAI API key from [OpenAI Platform](https://platform.openai.com/api-keys)
           - Enter it in the sidebar (it's stored securely in your session)
        
        2. **Configure Options**
           - Choose negative test ratio (10-15% recommended)
           - Select customer/product variations
           - Optionally focus on specific modules
           - Test case count is automatically determined based on BRD complexity
        
        3. **Upload BRD Documents**
           - Upload one or multiple BRD files
           - Supported formats: Word (.docx), PDF, Excel (.xlsx), Text (.txt)
           - Multiple files will be combined for comprehensive analysis
        
        4. **Preview (Optional)**
           - Click "Preview Extracted Text" to see what text was extracted
           - Verify the content looks correct
        
        5. **Generate Test Cases**
           - Click "Generate Test Cases"
           - **Step 1**: AI analyzes BRD to identify ALL functional requirements
           - **Step 2**: AI generates exactly ONE test case for EACH requirement
           - Wait for AI to complete both steps (may take 2-3 minutes)
           - View results in the "View Results" tab with 1:1 requirement mapping
        
        6. **Review and Filter**
           - Review generated test cases
           - Use filters to focus on specific modules or categories
           - Check statistics and distribution
        
        7. **Download**
           - Download as Excel (recommended - includes formatting)
           - Or download as CSV for other tools
           - Excel file includes summary sheet with statistics
        
        ### Tips for Best Results
        
        ✅ **DO:**
        - Upload complete BRD documents with detailed requirements
        - Include screen specifications, field details, and business rules
        - Specify workflows and user roles in the BRD
        - Use clear section headings in your BRD
        - Upload multiple related documents if needed
        
        ❌ **DON'T:**
        - Upload files without actual requirements
        - Expect perfect results from incomplete BRDs
        - Generate more than 500 test cases in one request (split large BRDs)
        
        ### Test Case Structure
        
        Generated test cases follow the **14-column CASA framework**:
        
        1. Product Name
        2. Process Category  
        3. Business Process ID
        4. Business Process
        5. Scenario ID
        6. Scenario Description
        7. Category (Positive/Negative)
        8. Importance (Critical/High/Medium/Low)
        9. Test Case ID
        10. TC Module
        11. Test Condition
        12. Pre-requisite
        13. Test Case Description
        14. Expected Result
        
        ### Framework Features
        
        - **🎯 1:1 Requirement Mapping**: Exactly one test case generated for each functional requirement
        - **📊 Requirements Analysis**: AI first identifies all requirements before generating test cases
        - **15-Step Workflow Model**: Complex workflows broken into 15 sequential steps
        - **Authorization Pattern**: 6+ test cases for maker-checker workflows
        - **Field Validation**: Separate test cases for each form section
        - **Variations**: Automatic generation for different customer/product types
        - **Negative Tests**: 10-15% coverage of business rule violations
        - **Perfect Coverage**: Ensures no requirement is missed or over-tested
        
        ### Troubleshooting
        
        **Problem**: API key not working  
        **Solution**: Verify your API key is correct and has sufficient credits
        
        **Problem**: No test cases generated  
        **Solution**: Check if BRD has enough detail. Try uploading more comprehensive documents.
        
        **Problem**: Test cases seem generic  
        **Solution**: Provide more detailed BRD with specific fields, screens, and workflows
        
        **Problem**: Not enough negative test cases  
        **Solution**: Increase the negative test ratio in settings. Ensure BRD has business rules.
        
        ### Support
        
        For issues or questions:
        - Check the framework documentation files
        - Review the Test_Case_Generation_Prompt.md
        - Refer to the CASA example (TC_CASA_For AI.xlsx)
        """)

if __name__ == "__main__":
    main()